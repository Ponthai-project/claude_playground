"""
AWS AI Practitioner (AIF-C01) 学習ロードマップ生成スクリプト
高坂昌信 編成
"""

import openpyxl
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
from openpyxl.styles.numbers import FORMAT_TEXT
from openpyxl.formatting.rule import ColorScaleRule
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ========================
#  色定義
# ========================
COLORS = {
    "header_dark":  "1F3864",  # 濃紺
    "header_mid":   "2E75B6",  # 青
    "phase1":       "D6E4F0",  # 薄青
    "phase2":       "D5E8D4",  # 薄緑
    "phase3":       "FFF2CC",  # 薄黄
    "phase4":       "FCE4D6",  # 薄オレンジ
    "phase5":       "E2EFDA",  # 薄緑2
    "white":        "FFFFFF",
    "light_gray":   "F2F2F2",
    "border":       "BFBFBF",
    "accent_red":   "C00000",
    "accent_green": "375623",
    "text_dark":    "1A1A1A",
}

PHASE_COLORS = ["D6E4F0", "D5E8D4", "FFF2CC", "FCE4D6", "E2EFDA"]
PHASE_HEADER_COLORS = ["2E75B6", "70AD47", "FFC000", "ED7D31", "70AD47"]


def make_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)


def make_font(bold=False, size=11, color="1A1A1A", name="Yu Gothic"):
    return Font(bold=bold, size=size, color=color, name=name)


def make_border(style="thin"):
    s = Side(border_style=style, color=COLORS["border"])
    return Border(left=s, right=s, top=s, bottom=s)


def make_center():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)


def make_left():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)


# ========================
#  Excel 作成
# ========================
def create_excel():
    wb = openpyxl.Workbook()

    sheet_overview(wb)
    sheet_weekly(wb)
    sheet_resources(wb)

    out_path = os.path.join(OUTPUT_DIR, "AWS_AI_Practitioner_学習ロードマップ.xlsx")
    wb.save(out_path)
    print(f"[Excel] 保存完了: {out_path}")
    return out_path


# ---------- Sheet 1: 全体ロードマップ ----------
def sheet_overview(wb):
    ws = wb.active
    ws.title = "全体ロードマップ"
    ws.sheet_view.showGridLines = False

    # タイトル
    ws.merge_cells("B2:K2")
    c = ws["B2"]
    c.value = "AWS Certified AI Practitioner (AIF-C01)　学習ロードマップ（2ヶ月）"
    c.font = make_font(bold=True, size=16, color="FFFFFF")
    c.fill = make_fill(COLORS["header_dark"])
    c.alignment = make_center()
    ws.row_dimensions[2].height = 36

    # サブタイトル
    ws.merge_cells("B3:K3")
    c = ws["B3"]
    c.value = "試験日標: 2026年7月初旬 ／ 推奨学習時間: 週10〜15時間 ／ 合格ライン: 700/1000点"
    c.font = make_font(size=10, color="FFFFFF")
    c.fill = make_fill(COLORS["header_mid"])
    c.alignment = make_center()
    ws.row_dimensions[3].height = 22

    # フェーズ定義
    phases = [
        {
            "no": 1,
            "name": "AI/ML 基礎固め",
            "period": "Week 1〜2（5/5〜5/18）",
            "hours": "約20時間",
            "goal": "AI・MLの根本概念を理解し、試験の全体像を把握する",
            "topics": [
                "・AI / ML / Deep Learning の定義と違い",
                "・教師あり学習 / 教師なし学習 / 強化学習",
                "・MLワークフロー（収集→前処理→訓練→評価→デプロイ）",
                "・過学習・アンダーフィット・評価指標（Accuracy, F1, AUC）",
                "・AWSにおけるAI/MLサービスの全体マップ",
            ],
            "deliverable": "AI/ML基礎の自己チェックシートを完成",
            "color": PHASE_COLORS[0],
            "hcolor": PHASE_HEADER_COLORS[0],
        },
        {
            "no": 2,
            "name": "AWS AI/ML マネージドサービス",
            "period": "Week 3〜4（5/19〜6/1）",
            "hours": "約25時間",
            "goal": "AWSが提供するAI/MLサービスの機能・ユースケースを習得",
            "topics": [
                "・Amazon SageMaker（Studio / Autopilot / Pipelines / Endpoints）",
                "・Amazon Rekognition / Comprehend / Textract / Transcribe",
                "・Amazon Polly / Translate / Forecast / Personalize",
                "・Amazon Kendra / Lex / Connect との連携",
                "・各サービスのユースケース・制限・料金体系",
            ],
            "deliverable": "サービス比較一覧表（自作）の完成",
            "color": PHASE_COLORS[1],
            "hcolor": PHASE_HEADER_COLORS[1],
        },
        {
            "no": 3,
            "name": "生成AI・基盤モデル（GenAI）",
            "period": "Week 5〜6（6/2〜6/15）",
            "hours": "約25時間",
            "goal": "生成AI・LLM・Amazon Bedrockを中心とした最新知識を習得",
            "topics": [
                "・生成AI の基本概念（Transformer / Attention / Tokenization）",
                "・Large Language Models（LLM）の仕組みと限界",
                "・Amazon Bedrock（基盤モデル選択 / Agents / Knowledge Bases）",
                "・プロンプトエンジニアリング（Zero-shot / Few-shot / CoT / RAG）",
                "・Fine-tuning vs RAG の使い分け・コスト比較",
            ],
            "deliverable": "Bedrock プロンプト比較メモの完成",
            "color": PHASE_COLORS[2],
            "hcolor": PHASE_HEADER_COLORS[2],
        },
        {
            "no": 4,
            "name": "責任あるAI・セキュリティ・ガバナンス",
            "period": "Week 7（6/16〜6/22）",
            "hours": "約15時間",
            "goal": "試験の14%×2領域（Responsible AI + Security）を完全カバー",
            "topics": [
                "・AIの倫理・公平性・バイアス・透明性・説明可能性（XAI）",
                "・AWS責任あるAIの6原則",
                "・データプライバシー / GDPR / 個人情報保護の観点",
                "・IAM / VPC / KMS によるAIワークロードのセキュリティ",
                "・AWS Audit Manager / AWS Config によるコンプライアンス管理",
            ],
            "deliverable": "Responsible AI チェックリストの作成",
            "color": PHASE_COLORS[3],
            "hcolor": PHASE_HEADER_COLORS[3],
        },
        {
            "no": 5,
            "name": "総仕上げ・模擬試験",
            "period": "Week 8〜9（6/23〜7/5）",
            "hours": "約20時間",
            "goal": "模擬試験で700点以上を安定させ、本番試験に臨む",
            "topics": [
                "・AWS公式 模擬試験（Skill Builder 無料版）× 2回",
                "・Whizlabs / ExamTopics 演習問題 × 100問以上",
                "・弱点ドメインの集中復習（前フェーズの見直し）",
                "・試験当日の注意事項・ID・会場確認（またはオンライン受験設定）",
                "・最終チェック：全サービス名と主要ユースケースの暗記",
            ],
            "deliverable": "模擬試験 700点突破 → 本番受験",
            "color": PHASE_COLORS[4],
            "hcolor": PHASE_HEADER_COLORS[4],
        },
    ]

    # フェーズヘッダー行
    row = 5
    ws.row_dimensions[row].height = 20

    col_map = {
        "B": "フェーズ", "C": "期間", "D": "目標時間",
        "E": "学習目標", "F": "主要トピック（抜粋）", "G": "達成の目安",
    }
    header_cols = ["B", "C", "D", "E", "F", "G"]
    col_widths = [16, 24, 14, 36, 50, 30]

    for i, (col, label) in enumerate(zip(header_cols, col_map.values())):
        c = ws[f"{col}{row}"]
        c.value = label
        c.font = make_font(bold=True, size=11, color="FFFFFF")
        c.fill = make_fill(COLORS["header_dark"])
        c.alignment = make_center()
        c.border = make_border()
        ws.column_dimensions[col].width = col_widths[i]

    for ph in phases:
        row += 1
        ws.row_dimensions[row].height = 90

        cells = {
            "B": f"Phase {ph['no']}\n{ph['name']}",
            "C": ph["period"],
            "D": ph["hours"],
            "E": ph["goal"],
            "F": "\n".join(ph["topics"]),
            "G": ph["deliverable"],
        }
        for col, val in cells.items():
            c = ws[f"{col}{row}"]
            c.value = val
            c.fill = make_fill(ph["color"])
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            c.border = make_border()
            c.font = make_font(size=10)
            if col == "B":
                c.font = make_font(bold=True, size=10, color=ph["hcolor"])
                c.alignment = make_center()

    # ガントチャート簡易版
    row += 2
    ws.merge_cells(f"B{row}:K{row}")
    c = ws[f"B{row}"]
    c.value = "▼ 月別 ガントチャート"
    c.font = make_font(bold=True, size=12, color="FFFFFF")
    c.fill = make_fill(COLORS["header_dark"])
    c.alignment = make_center()
    ws.row_dimensions[row].height = 24

    row += 1
    ws.row_dimensions[row].height = 18
    gantt_months = ["", "5月上", "5月中", "5月下", "6月上", "6月中", "6月下", "7月上"]
    gantt_cols = ["B", "C", "D", "E", "F", "G", "H", "I"]
    for col, label in zip(gantt_cols, gantt_months):
        c = ws[f"{col}{row}"]
        c.value = label
        c.font = make_font(bold=True, size=10, color="FFFFFF")
        c.fill = make_fill(COLORS["header_mid"])
        c.alignment = make_center()
        c.border = make_border()

    gantt_data = [
        ("Phase 1: AI/ML基礎",    ["■", "■", "", "", "", "", ""]),
        ("Phase 2: AWSサービス",  ["", "", "■", "■", "", "", ""]),
        ("Phase 3: 生成AI",       ["", "", "", "", "■", "■", ""]),
        ("Phase 4: 責任AI",       ["", "", "", "", "", "■", ""]),
        ("Phase 5: 模擬試験",     ["", "", "", "", "", "■", "■"]),
    ]
    phase_fill_colors = PHASE_COLORS

    for idx, (label, marks) in enumerate(gantt_data):
        row += 1
        ws.row_dimensions[row].height = 18
        c = ws[f"B{row}"]
        c.value = label
        c.font = make_font(bold=True, size=10, color=PHASE_HEADER_COLORS[idx])
        c.fill = make_fill(COLORS["light_gray"])
        c.border = make_border()
        c.alignment = make_left()

        for col, mark in zip(gantt_cols[1:], marks):
            c = ws[f"{col}{row}"]
            if mark == "■":
                c.value = "●"
                c.fill = make_fill(PHASE_COLORS[idx])
                c.font = make_font(bold=True, size=12, color=PHASE_HEADER_COLORS[idx])
            else:
                c.fill = make_fill(COLORS["white"])
            c.alignment = make_center()
            c.border = make_border()

    ws.column_dimensions["A"].width = 2


# ---------- Sheet 2: 週別詳細計画 ----------
def sheet_weekly(wb):
    ws = wb.create_sheet("週別詳細計画")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 2

    ws.merge_cells("B2:J2")
    c = ws["B2"]
    c.value = "週別 詳細学習計画"
    c.font = make_font(bold=True, size=14, color="FFFFFF")
    c.fill = make_fill(COLORS["header_dark"])
    c.alignment = make_center()
    ws.row_dimensions[2].height = 30

    headers = ["週", "期間", "フェーズ", "月曜〜水曜（平日前半）", "木曜〜金曜（平日後半）", "土日（週末）", "推奨サイト/教材", "週次チェック"]
    col_letters = ["B", "C", "D", "E", "F", "G", "H", "I"]
    col_widths_w = [8, 22, 20, 38, 38, 38, 32, 28]

    row = 3
    ws.row_dimensions[row].height = 20
    for col, label, w in zip(col_letters, headers, col_widths_w):
        c = ws[f"{col}{row}"]
        c.value = label
        c.font = make_font(bold=True, size=10, color="FFFFFF")
        c.fill = make_fill(COLORS["header_mid"])
        c.alignment = make_center()
        c.border = make_border()
        ws.column_dimensions[col].width = w

    weekly = [
        {
            "week": "Week 1", "period": "5/5（月）〜5/11（日）",
            "phase": "Phase 1\nAI/ML基礎①",
            "mon_wed": "・AIとは何か / MLとの違い\n・教師あり / 教師なし / 強化学習を図解で理解\n（Skill Builder: AI/ML基礎コース視聴）",
            "thu_fri": "・MLワークフロー全体像（データ→モデル→推論）\n・主要アルゴリズム概要（回帰・決定木・SVM・NN）\n（AWS公式ドキュメント読み込み）",
            "weekend": "・深層学習基礎（CNN/RNN/Transformer概要）\n・評価指標（Precision/Recall/F1/AUC）の理解\n・Week1 内容をノートにまとめる（1時間）",
            "resource": "AWS Skill Builder\n「Exploring AI/ML」コース\nBlack Belt: ML基礎",
            "check": "□ AI/ML/DLの違いを説明できる\n□ 3種の学習方法を言える\n□ 評価指標を使い分けられる",
            "phasenum": 0,
        },
        {
            "week": "Week 2", "period": "5/12（月）〜5/18（日）",
            "phase": "Phase 1\nAI/ML基礎②",
            "mon_wed": "・過学習・バイアス・バリアンスの概念\n・正則化（L1/L2）とドロップアウト\n・特徴量エンジニアリングの基本",
            "thu_fri": "・AWSのAI/MLサービス全体マップを整理\n・SageMakerの概要と主要機能の把握\n（公式ドキュメント・BlackBelt視聴）",
            "weekend": "・Phase1 全体復習＋自己チェックシート作成\n・模擬問題（Skill Builder 無料サンプル10問）\n・次週の予習（SageMaker概要）",
            "resource": "AWS Skill Builder\n「Exam Prep: AI Practitioner」\n（無料版）第1章",
            "check": "□ 過学習の対策を3つ言える\n□ AWSのAI系サービス10個挙げられる\n□ SageMakerの役割を説明できる",
            "phasenum": 0,
        },
        {
            "week": "Week 3", "period": "5/19（月）〜5/25（日）",
            "phase": "Phase 2\nAWSサービス①",
            "mon_wed": "・Amazon SageMaker 詳細\n　（Studio / Autopilot / Clarify / Debugger）\n・SageMaker Pipelines / Model Registry",
            "thu_fri": "・Amazon Rekognition（画像・動画分析）\n・Amazon Comprehend（自然言語処理）\n・Amazon Textract（OCR/文書解析）",
            "weekend": "・Amazon Transcribe / Polly / Translate 復習\n・ユースケース問題演習（各サービス）\n・サービス比較表の作成開始",
            "resource": "AWS公式ドキュメント\nSageMaker BlackBelt\nAWS re:Invent SageMaker動画",
            "check": "□ SageMaker各機能の使い分けを説明できる\n□ Rekognition/Comprehendの違いを言える\n□ ユースケース3問正解できる",
            "phasenum": 1,
        },
        {
            "week": "Week 4", "period": "5/26（月）〜6/1（日）",
            "phase": "Phase 2\nAWSサービス②",
            "mon_wed": "・Amazon Forecast（時系列予測）\n・Amazon Personalize（レコメンドエンジン）\n・Amazon Kendra（インテリジェント検索）",
            "thu_fri": "・Amazon Lex（チャットボット）\n・AWS Panorama / AWS DeepRacer 概要\n・サービス選択問題の解き方のコツ習得",
            "weekend": "・Phase2 全体復習\n・サービス比較表を完成させる\n・Whizlabs演習問題 30問（Phase1-2範囲）",
            "resource": "Whizlabs AWS AI Practitioner\nAWS FAQ各サービスページ\nExamTopics AIF-C01",
            "check": "□ Forecastとpersonalizeの違いを言える\n□ Kendra vs. OpenSearchの使い分け\n□ 模擬問題30問中20問以上正解",
            "phasenum": 1,
        },
        {
            "week": "Week 5", "period": "6/2（月）〜6/8（日）",
            "phase": "Phase 3\n生成AI①",
            "mon_wed": "・生成AIとは / 従来MLとの違い\n・Transformer / Self-Attention の仕組み\n・LLM（GPT系）の訓練と推論の流れ",
            "thu_fri": "・Amazon Bedrock の概要\n・利用可能な基盤モデル（Claude / Titan / Llama等）\n・Bedrockのユースケースとコスト構造",
            "weekend": "・プロンプトエンジニアリング実践\n　（Zero-shot / Few-shot / CoT / System Prompt）\n・Bedrock Playground で実際に試す（無料枠）",
            "resource": "AWS Skill Builder\n「Generative AI with LLMs」\nAmazon Bedrock ドキュメント",
            "check": "□ Transformerの概念を説明できる\n□ Bedrockで使える基盤モデルを4つ言える\n□ プロンプト技法の違いを説明できる",
            "phasenum": 2,
        },
        {
            "week": "Week 6", "period": "6/9（月）〜6/15（日）",
            "phase": "Phase 3\n生成AI②",
            "mon_wed": "・RAG（検索拡張生成）の仕組みと実装\n・Bedrock Knowledge Bases の設定方法\n・ベクトルDB（OpenSearch Serverless）連携",
            "thu_fri": "・Fine-tuning の概念と適用シナリオ\n・RLHF（人間フィードバックによる強化学習）\n・Bedrock Agents / Function Calling",
            "weekend": "・Phase3 全体復習\n・Bedrockプロンプト比較メモ完成\n・模擬試験（Skill Builder 公式 Exam Prep）30問",
            "resource": "AWS Bedrock ユーザーガイド\nAWS公式Blog: RAG実装例\nSkill Builder Exam Prep 第3章",
            "check": "□ RAGとFine-tuningの使い分けを説明できる\n□ Bedrock Agentsの役割を言える\n□ 模擬問題 30問中22問以上正解",
            "phasenum": 2,
        },
        {
            "week": "Week 7", "period": "6/16（月）〜6/22（日）",
            "phase": "Phase 4\n責任AI・セキュリティ",
            "mon_wed": "・AI倫理の6原則（公平性/透明性/説明可能性等）\n・AWS Responsible AI フレームワーク\n・SageMaker Clarify によるバイアス検出",
            "thu_fri": "・AIワークロードのセキュリティ設計\n　（IAM / VPC / KMS / PrivateLink）\n・AWS Artifact / Audit Manager / Config",
            "weekend": "・Responsible AI チェックリスト作成\n・Phase4 まとめ復習\n・Whizlabs演習 40問（全フェーズ混合）",
            "resource": "AWS Responsible AI ページ\nAWS Security Whitepaper\nNIST AI RMF（参考）",
            "check": "□ AI倫理の6原則を列挙できる\n□ AIセキュリティのAWSサービスを3つ挙げられる\n□ 演習40問中30問以上正解",
            "phasenum": 3,
        },
        {
            "week": "Week 8", "period": "6/23（月）〜6/29（日）",
            "phase": "Phase 5\n模擬試験①",
            "mon_wed": "・Skill Builder 公式模擬試験（65問）を受験\n・結果分析：ドメイン別スコアを記録\n・スコアが低いドメインを特定",
            "thu_fri": "・弱点ドメインの集中復習（2〜3時間）\n・ExamTopics の最新問題 30問演習\n・間違えた問題をすべて解説読み直し",
            "weekend": "・2回目の模擬試験（別の問題セット）受験\n・700点相当（約72%正解）を目標\n・試験登録・会場（またはオンライン）確認",
            "resource": "AWS Skill Builder 模擬試験\nExamTopics AIF-C01\nWhizlabs Practice Test",
            "check": "□ 1回目模擬試験 65%以上\n□ 2回目模擬試験 70%以上\n□ 試験日の登録完了",
            "phasenum": 4,
        },
        {
            "week": "Week 9", "period": "6/30（月）〜7/5（土）",
            "phase": "Phase 5\n最終仕上げ・本番",
            "mon_wed": "・全サービス名と主要ユースケース最終暗記\n・弱点分野の絞り込み復習（2時間）\n・3回目模擬試験（目標：75%以上）",
            "thu_fri": "・試験前日：軽い復習のみ（新規学習NG）\n・試験環境チェック（オンラインの場合）\n・IDと受験票の確認・体調管理",
            "weekend": "★ 本番試験受験 ★\n試験後：結果確認（即時スコア表示）\n合格後：デジタルバッジ取得申請",
            "resource": "AWS Certification ポータル\n（試験スケジュール・ID確認）",
            "check": "□ 3回目模擬試験 75%以上\n□ 本番試験 受験完了\n□ 合格（700点以上）",
            "phasenum": 4,
        },
    ]

    for wk in weekly:
        row += 1
        ws.row_dimensions[row].height = 95
        pidx = wk["phasenum"]
        cells = {
            "B": wk["week"],
            "C": wk["period"],
            "D": wk["phase"],
            "E": wk["mon_wed"],
            "F": wk["thu_fri"],
            "G": wk["weekend"],
            "H": wk["resource"],
            "I": wk["check"],
        }
        for col, val in cells.items():
            c = ws[f"{col}{row}"]
            c.value = val
            c.fill = make_fill(PHASE_COLORS[pidx])
            c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            c.border = make_border()
            c.font = make_font(size=9)
            if col in ("B", "C", "D"):
                c.font = make_font(bold=True, size=9, color=PHASE_HEADER_COLORS[pidx])
                c.alignment = make_center()


# ---------- Sheet 3: 学習リソース ----------
def sheet_resources(wb):
    ws = wb.create_sheet("学習リソース一覧")
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 2

    ws.merge_cells("B2:H2")
    c = ws["B2"]
    c.value = "学習リソース一覧 ― AWS AI Practitioner (AIF-C01)"
    c.font = make_font(bold=True, size=14, color="FFFFFF")
    c.fill = make_fill(COLORS["header_dark"])
    c.alignment = make_center()
    ws.row_dimensions[2].height = 30

    headers_r = ["カテゴリ", "教材名", "URL / 入手先", "費用", "優先度", "用途・特徴"]
    col_letters_r = ["B", "C", "D", "E", "F", "G"]
    col_widths_r = [18, 36, 52, 14, 12, 44]

    row = 3
    for col, label, w in zip(col_letters_r, headers_r, col_widths_r):
        c = ws[f"{col}{row}"]
        c.value = label
        c.font = make_font(bold=True, size=10, color="FFFFFF")
        c.fill = make_fill(COLORS["header_mid"])
        c.alignment = make_center()
        c.border = make_border()
        ws.column_dimensions[col].width = w

    resources = [
        # カテゴリ, 名称, URL, 費用, 優先度, 用途
        ("【公式】試験情報",
         "AWS Certified AI Practitioner 公式ページ",
         "https://aws.amazon.com/jp/certification/certified-ai-practitioner/",
         "無料", "★★★", "試験概要・ドメイン・登録リンク。まず最初に読む"),
        ("【公式】試験情報",
         "試験ガイド (PDF / 日本語)",
         "https://d1.awsstatic.com/ja_JP/training-and-certification/docs-ai-practitioner/AWS-Certified-AI-Practitioner_Exam-Guide.pdf",
         "無料", "★★★", "出題範囲・比率が明記。印刷して手元に置く"),
        ("【公式】学習コース",
         "AWS Skill Builder: Exam Prep (AIF-C01)",
         "https://skillbuilder.aws/",
         "無料（一部有料）", "★★★", "AWS公式の試験対策コース。無料版でも基本カバー"),
        ("【公式】学習コース",
         "AWS Skill Builder: Generative AI for Execs & Decision Makers",
         "https://skillbuilder.aws/",
         "無料", "★★★", "生成AIの基礎をビジネス視点で学習。AIF試験の核心"),
        ("【公式】学習コース",
         "AWS Skill Builder: Amazon Bedrock Getting Started",
         "https://skillbuilder.aws/",
         "無料", "★★★", "Bedrock入門コース。Week5-6の必修"),
        ("【公式】模擬試験",
         "AWS Skill Builder: Official Practice Question Set (AIF-C01)",
         "https://skillbuilder.aws/",
         "無料", "★★★", "公式の練習問題セット（20問）。問題の傾向把握に最適"),
        ("【公式】ドキュメント",
         "Amazon SageMaker ドキュメント",
         "https://docs.aws.amazon.com/ja_jp/sagemaker/",
         "無料", "★★★", "SageMaker各機能のリファレンス"),
        ("【公式】ドキュメント",
         "Amazon Bedrock ユーザーガイド",
         "https://docs.aws.amazon.com/ja_jp/bedrock/",
         "無料", "★★★", "Bedrock全機能の公式ガイド"),
        ("【公式】ドキュメント",
         "AWS Responsible AI",
         "https://aws.amazon.com/jp/machine-learning/responsible-machine-learning/",
         "無料", "★★☆", "責任あるAIの公式説明。試験の14%をカバー"),
        ("【外部】演習問題",
         "ExamTopics: AWS AI Practitioner",
         "https://www.examtopics.com/exams/amazon/aws-certified-ai-practitioner/",
         "無料（一部有料）", "★★★", "実際の試験に近い問題集。コミュニティ解説付き"),
        ("【外部】演習問題",
         "Whizlabs AWS Certified AI Practitioner",
         "https://www.whizlabs.com/aws-certified-ai-practitioner/",
         "約$17〜29", "★★☆", "本番に近い品質の模擬試験。Udemy セールで安くなることも"),
        ("【外部】動画講座",
         "Udemy: AWS AI Practitioner 対策コース（日本語）",
         "https://www.udemy.com/（「AWS AI Practitioner」で検索）",
         "約¥2,000（セール時）", "★★☆", "日本語で体系的に学べる。セール時（¥1,500〜）に購入推奨"),
        ("【外部】動画講座",
         "YouTube: FreeCodeCamp AWS AI Practitioner",
         "https://www.youtube.com/（「AWS AI Practitioner FreeCodeCamp」で検索）",
         "無料", "★★☆", "英語だが無料で10時間以上の講義動画"),
        ("【参考】AI/ML基礎",
         "Coursera: Machine Learning Specialization (Andrew Ng)",
         "https://www.coursera.org/specializations/machine-learning-introduction",
         "無料（聴講のみ）", "★☆☆", "ML基礎の深掘りに。試験には少し深すぎるが概念理解に有用"),
        ("【AWS】BlackBelt",
         "AWS Black Belt: Amazon SageMaker",
         "https://www.youtube.com/（「AWS Black Belt SageMaker」で検索）",
         "無料", "★★☆", "AWS公式の詳細解説。各サービスのBlackBeltを1本ずつ視聴"),
        ("【試験登録】",
         "AWS Certification ポータル（試験申込み）",
         "https://aws.amazon.com/jp/certification/",
         "USD $150", "★★★", "Pearson VUE または PSI で受験予約。受験料の割引バウチャーはAWS re/Startや無料コース完了後に付与される場合あり"),
    ]

    cat_colors = {
        "【公式】試験情報": "D6E4F0",
        "【公式】学習コース": "D5E8D4",
        "【公式】模擬試験": "D5E8D4",
        "【公式】ドキュメント": "D5E8D4",
        "【外部】演習問題": "FFF2CC",
        "【外部】動画講座": "FFF2CC",
        "【参考】AI/ML基礎": "F2F2F2",
        "【AWS】BlackBelt": "FCE4D6",
        "【試験登録】": "E2EFDA",
    }

    for res in resources:
        row += 1
        ws.row_dimensions[row].height = 40
        cat, name, url, cost, priority, desc = res
        fill_color = cat_colors.get(cat, "FFFFFF")
        vals = [cat, name, url, cost, priority, desc]
        for col, val in zip(col_letters_r, vals):
            c = ws[f"{col}{row}"]
            c.value = val
            c.fill = make_fill(fill_color)
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            c.border = make_border()
            c.font = make_font(size=9)
            if col == "F" and val == "★★★":
                c.font = make_font(bold=True, size=10, color=COLORS["accent_red"])
                c.alignment = make_center()
            elif col == "F":
                c.font = make_font(size=10, color="595959")
                c.alignment = make_center()
            elif col == "D":
                c.font = make_font(size=9, color="0563C1")


# ========================
#  Word 作成
# ========================
def create_word():
    doc = Document()

    # ページマージン設定
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # タイトル
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("AWS Certified AI Practitioner (AIF-C01)")
    run.font.name = "Yu Gothic"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("学習ロードマップ（2ヶ月完成プラン）")
    run2.font.name = "Yu Gothic"
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = info_para.add_run("学習期間: 2026年5月5日〜7月5日 ／ 合格目標点: 700/1000 ／ 推奨学習時間: 週10〜15時間")
    run3.font.name = "Yu Gothic"
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # === 試験概要 ===
    _add_heading(doc, "1. 試験概要", level=1)

    overview_data = [
        ("試験コード", "AIF-C01"),
        ("試験名（日本語）", "AWS Certified AI Practitioner"),
        ("出題数", "65問（スコア対象）＋ 15問（未公開の試験問題）"),
        ("試験時間", "120分"),
        ("受験料", "USD $150（日本円で約22,500円）"),
        ("合格ライン", "700 / 1000点"),
        ("試験形式", "多肢選択・複数選択"),
        ("受験方法", "Pearson VUE（テストセンター）または自宅オンライン"),
        ("有効期限", "3年間"),
    ]

    table = doc.add_table(rows=len(overview_data), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(overview_data):
        row_cells = table.rows[i].cells
        row_cells[0].text = k
        row_cells[1].text = v
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Yu Gothic"
                    run.font.size = Pt(10)
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # === 出題ドメイン ===
    _add_heading(doc, "2. 出題ドメインと配点", level=1)

    domain_data = [
        ("ドメイン", "内容", "比率"),
        ("Domain 1", "AI/MLの基礎（Foundation of AI and ML）", "20%"),
        ("Domain 2", "生成AIの基礎（Foundation of Generative AI）", "24%"),
        ("Domain 3", "基盤モデルの応用（Applications of Foundation Models）", "28%"),
        ("Domain 4", "責任あるAI（Guidelines for Responsible AI）", "14%"),
        ("Domain 5", "AIソリューションのセキュリティ・ガバナンス", "14%"),
    ]

    table2 = doc.add_table(rows=len(domain_data), cols=3)
    table2.style = "Table Grid"
    for i, row_data in enumerate(domain_data):
        row_cells = table2.rows[i].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val
            for para in row_cells[j].paragraphs:
                for run in para.runs:
                    run.font.name = "Yu Gothic"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    doc.add_paragraph()

    # === フェーズ別ロードマップ ===
    _add_heading(doc, "3. フェーズ別 学習ロードマップ", level=1)

    phases_w = [
        {
            "no": 1, "name": "AI/ML 基礎固め",
            "period": "Week 1〜2（2026/5/5〜5/18）",
            "hours": "約20時間",
            "goal": "AI・MLの根本概念を理解し、試験の全体像を把握する",
            "items": [
                "AI / ML / Deep Learning の定義と違いを図で理解",
                "教師あり学習・教師なし学習・強化学習の3分類",
                "MLワークフロー（データ収集→前処理→訓練→評価→デプロイ）",
                "主要アルゴリズム概要（回帰・決定木・SVM・ニューラルネット）",
                "評価指標：Accuracy / Precision / Recall / F1 / AUC-ROC",
                "深層学習基礎：CNN（画像）/ RNN（時系列）/ Transformer（言語）",
                "AWSのAI/MLサービス全体マップを俯瞰",
            ],
            "sites": [
                ("AWS Skill Builder 「Exploring AI/ML」コース", "https://skillbuilder.aws/"),
                ("AWS Certified AI Practitioner 公式ページ", "https://aws.amazon.com/jp/certification/certified-ai-practitioner/"),
                ("試験ガイド (PDF)", "https://d1.awsstatic.com/ja_JP/training-and-certification/docs-ai-practitioner/AWS-Certified-AI-Practitioner_Exam-Guide.pdf"),
            ],
        },
        {
            "no": 2, "name": "AWS AI/ML マネージドサービス",
            "period": "Week 3〜4（2026/5/19〜6/1）",
            "hours": "約25時間",
            "goal": "AWSが提供するAI/MLマネージドサービスの機能・ユースケースを完全習得",
            "items": [
                "Amazon SageMaker: Studio / Autopilot / Clarify / Debugger / Pipelines",
                "Amazon Rekognition: 画像・動画認識（顔認証・物体検出）",
                "Amazon Comprehend: テキスト分析（感情分析・エンティティ抽出）",
                "Amazon Textract: OCR・帳票解析",
                "Amazon Transcribe: 音声→テキスト変換",
                "Amazon Polly: テキスト→音声変換",
                "Amazon Translate: 機械翻訳",
                "Amazon Forecast: 時系列予測",
                "Amazon Personalize: レコメンデーションエンジン",
                "Amazon Kendra: インテリジェント検索 / Amazon Lex: チャットボット",
            ],
            "sites": [
                ("Amazon SageMaker ドキュメント", "https://docs.aws.amazon.com/ja_jp/sagemaker/"),
                ("AWS Black Belt SageMaker（YouTube）", "https://www.youtube.com/"),
                ("ExamTopics AIF-C01", "https://www.examtopics.com/exams/amazon/aws-certified-ai-practitioner/"),
            ],
        },
        {
            "no": 3, "name": "生成AI・基盤モデル（GenAI）",
            "period": "Week 5〜6（2026/6/2〜6/15）",
            "hours": "約25時間",
            "goal": "生成AI・LLM・Amazon Bedrockを中心とした最新技術を深く理解する",
            "items": [
                "生成AIの基本概念（Transformer / Attention / Tokenization）",
                "大規模言語モデル（LLM）の訓練・推論・チューニングの流れ",
                "Amazon Bedrock: 基盤モデル選択（Claude / Titan / Llama / Mistral等）",
                "Bedrock Knowledge Bases によるRAG実装",
                "Bedrock Agents / Function Calling による自律エージェント構築",
                "プロンプトエンジニアリング: Zero-shot / Few-shot / CoT / System Prompt",
                "RAG（検索拡張生成）vs Fine-tuning の使い分け・コスト比較",
                "ハルシネーション（幻覚）の原因と対策",
                "ベクトルDB: OpenSearch Serverless / Aurora pgvector",
            ],
            "sites": [
                ("Amazon Bedrock ユーザーガイド", "https://docs.aws.amazon.com/ja_jp/bedrock/"),
                ("AWS Skill Builder: Generative AI Fundamentals", "https://skillbuilder.aws/"),
                ("AWS Skill Builder: Amazon Bedrock Getting Started", "https://skillbuilder.aws/"),
            ],
        },
        {
            "no": 4, "name": "責任あるAI・セキュリティ・ガバナンス",
            "period": "Week 7（2026/6/16〜6/22）",
            "hours": "約15時間",
            "goal": "試験の28%を占める責任AI・セキュリティ領域を完全カバーする",
            "items": [
                "AI倫理の6原則: 公平性・透明性・プライバシー・セキュリティ・堅牢性・説明可能性",
                "AWS Responsible AI フレームワークと実装ガイドライン",
                "SageMaker Clarify によるバイアス検出と説明可能性（XAI）",
                "AIデータのプライバシー保護（GDPR / 個人情報保護）",
                "AIワークロードのセキュリティ: IAM / VPC / KMS / PrivateLink",
                "AWS Audit Manager / AWS Config によるコンプライアンス管理",
                "AWS Macie によるデータセキュリティ",
            ],
            "sites": [
                ("AWS Responsible AI 公式ページ", "https://aws.amazon.com/jp/machine-learning/responsible-machine-learning/"),
                ("AWS Security Whitepaper", "https://docs.aws.amazon.com/ja_jp/whitepapers/latest/introduction-aws-security/"),
                ("Whizlabs AWS AI Practitioner", "https://www.whizlabs.com/aws-certified-ai-practitioner/"),
            ],
        },
        {
            "no": 5, "name": "総仕上げ・模擬試験・本番受験",
            "period": "Week 8〜9（2026/6/23〜7/5）",
            "hours": "約20時間",
            "goal": "模擬試験で安定して700点（72%）以上取り、本番試験に合格する",
            "items": [
                "AWS Skill Builder 公式模擬試験（65問）を2回受験・分析",
                "ExamTopics で最新問題 100問以上演習",
                "Whizlabs 模擬試験セットを本番形式で受験",
                "弱点ドメインの集中復習（結果分析に基づき優先順位付け）",
                "全サービス名・主要ユースケース・制限の最終暗記",
                "試験前日: 新規学習禁止、軽い復習のみ（体調管理優先）",
                "本番受験 → 合格後デジタルバッジ取得申請",
            ],
            "sites": [
                ("AWS Skill Builder 公式模擬試験", "https://skillbuilder.aws/"),
                ("ExamTopics AIF-C01", "https://www.examtopics.com/exams/amazon/aws-certified-ai-practitioner/"),
                ("AWS Certification ポータル（試験登録）", "https://aws.amazon.com/jp/certification/"),
            ],
        },
    ]

    for ph in phases_w:
        _add_heading(doc, f"Phase {ph['no']}: {ph['name']}  ―  {ph['period']}  （{ph['hours']}）", level=2)

        goal_para = doc.add_paragraph()
        goal_run = goal_para.add_run(f"【目標】{ph['goal']}")
        goal_run.font.name = "Yu Gothic"
        goal_run.font.size = Pt(10)
        goal_run.font.bold = True
        goal_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        items_para = doc.add_paragraph()
        items_para.add_run("学習内容:").font.name = "Yu Gothic"
        for item in ph["items"]:
            p = doc.add_paragraph(f"  • {item}", style="List Bullet")
            for run in p.runs:
                run.font.name = "Yu Gothic"
                run.font.size = Pt(10)

        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run("参考サイト:")
        ref_run.font.name = "Yu Gothic"
        ref_run.font.size = Pt(10)
        ref_run.font.bold = True

        for site_name, url in ph["sites"]:
            p = doc.add_paragraph(f"  • {site_name}: {url}")
            for run in p.runs:
                run.font.name = "Yu Gothic"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

        doc.add_paragraph()

    # === 合格のコツ ===
    _add_heading(doc, "4. 合格のためのポイント", level=1)
    tips = [
        "【Domain 3が最重要】出題比率28%の「基盤モデルの応用」を最も重点的に学習。特にRAG・プロンプトエンジニアリング・Fine-tuningの使い分けを深く理解する。",
        "【サービスのユースケースを暗記】SageMaker / Bedrock / 各AIサービスの「何に使うか」を問う問題が多い。公式ドキュメントのユースケースセクションを必ず読む。",
        "【模擬試験は本番と同じ形式で】時間を測って65問を連続で解く練習を必ずする。疲労による後半の正答率低下を防ぐ。",
        "【日本語 or 英語の選択】日本語受験は追加30分付与（計150分）。英語に不安があれば日本語受験を推奨。ただし一部の翻訳が不自然な場合は英語版と照合する。",
        "【AWSの最新サービスに注意】Bedrockは進化が速い。試験は2024年秋リリースのAIF-C01バージョンに対応。試験ガイドのversion日付を必ず確認する。",
        "【Responsible AIは暗記勝負】倫理原則・AWSのフレームワーク・SageMaker Clarifyの機能は短時間で得点できる。Week7で一気に仕上げる。",
    ]
    for tip in tips:
        p = doc.add_paragraph(f"• {tip}")
        for run in p.runs:
            run.font.name = "Yu Gothic"
            run.font.size = Pt(10)

    doc.add_paragraph()

    # === 推奨学習スケジュール ===
    _add_heading(doc, "5. 推奨 日次学習スケジュール（参考）", level=1)
    schedule_data = [
        ("曜日", "推奨学習内容", "目安時間"),
        ("月曜", "新規インプット（動画講座 / ドキュメント読み）", "1.5〜2時間"),
        ("火曜", "前日の内容をノート整理・自分の言葉でまとめ", "1〜1.5時間"),
        ("水曜", "新規インプット続き（週のテーマを前半で完了）", "1.5〜2時間"),
        ("木曜", "演習問題（10〜20問）＋ 間違えた問題の解説確認", "1.5時間"),
        ("金曜", "サービス比較・暗記カード作成・弱点補強", "1時間"),
        ("土曜", "週まとめ復習 ＋ 模擬試験（最終週）", "2〜3時間"),
        ("日曜", "軽い復習 ＋ 翌週の予習（30分程度）＋ 休息", "0.5〜1時間"),
    ]

    table3 = doc.add_table(rows=len(schedule_data), cols=3)
    table3.style = "Table Grid"
    for i, row_data in enumerate(schedule_data):
        row_cells = table3.rows[i].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = val
            for para in row_cells[j].paragraphs:
                for run in para.runs:
                    run.font.name = "Yu Gothic"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

    out_path = os.path.join(OUTPUT_DIR, "AWS_AI_Practitioner_学習ロードマップ.docx")
    doc.save(out_path)
    print(f"[Word] 保存完了: {out_path}")
    return out_path


def _add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Yu Gothic"
    if level == 1:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)


if __name__ == "__main__":
    excel_path = create_excel()
    word_path = create_word()
    print("\n✅ 両ファイルの生成が完了しました。")
    print(f"  Excel: {excel_path}")
    print(f"  Word:  {word_path}")
