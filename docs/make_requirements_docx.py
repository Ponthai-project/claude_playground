from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ========== ページ設定 ==========
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ========== スタイル定義 ==========
def set_font(run, name="Yu Gothic", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 日本語フォント設定
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    # 左ボーダー風に■を付ける
    run_mark = p.add_run("■ ")
    set_font(run_mark, size=14, bold=True, color=(0x8b, 0x3a, 0x00))
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=(0x1a, 0x0a, 0x00))
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=(0x3d, 0x1a, 0x00))
    # 下線
    run.font.underline = True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("▷ " + text)
    set_font(run, size=10.5, bold=True, color=(0x5a, 0x2a, 0x00))
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_officer_comment(doc, speaker, comment, bg_hex, speaker_color, border_color):
    """武将コメントボックス"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

    # 枠線付きの段落（シェーディングで代用）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['left']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'thick')
        bdr.set(qn('w:sz'), '24')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), border_color)
        pBdr.append(bdr)
    pPr.append(pBdr)

    # 背景色
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)

    run_s = p.add_run(f"【{speaker}】\n")
    set_font(run_s, size=9, bold=True, color=tuple(int(speaker_color[i:i+2], 16) for i in (0,2,4)))
    run_c = p.add_run(comment)
    set_font(run_c, size=10)

def add_req_table(doc, headers, rows, col_widths=None):
    """要件テーブル"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '3D1A00')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=9.5, bold=True, color=(0xe8, 0xd5, 0xa3))

    # データ行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'FAF7F0' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9.5)

    # 列幅設定
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表の後に空行
    return table

# ========== 表紙 ==========
doc.add_paragraph()
doc.add_paragraph()

p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_cover.add_run("⚔")
set_font(r, size=48, color=(0x8b, 0x3a, 0x00))

p_label = doc.add_paragraph()
p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_label.add_run("Requirements Definition Document")
set_font(r, size=10, color=(0xc8, 0xa9, 0x6e))

doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("仮想プロレス団体経営アプリ")
set_font(r, size=28, bold=True, color=(0x1a, 0x0a, 0x00))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("要　件　定　義　書")
set_font(r, size=18, color=(0x8b, 0x3a, 0x00))

doc.add_paragraph()
doc.add_paragraph()

# 表紙メタ情報テーブル
meta_table = doc.add_table(rows=8, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("プロジェクト総大将", "殿下"),
    ("軍師（取りまとめ）", "武田信玄"),
    ("経営・内政担当",     "高坂昌信"),
    ("興行・戦闘担当",     "山県昌景"),
    ("品質・安定担当",     "馬場信春"),
    ("策略・物語担当",     "真田幸隆"),
    ("版　　　数",         "v1.0（初稿）"),
    ("作成日",             "2026年4月12日"),
]
for i, (k, v) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    set_cell_bg(c0, '3D1A00')
    set_cell_bg(c1, 'FAF7F0')
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = p0.add_run(k)
    set_font(r0, size=10, bold=True, color=(0xe8, 0xd5, 0xa3))
    r1 = c1.paragraphs[0].add_run(v)
    set_font(r1, size=10, color=(0x1a, 0x0a, 0x00))

doc.add_page_break()

# ========== 1. プロジェクト概要 ==========
add_heading1(doc, "1. プロジェクト概要")
add_officer_comment(doc,
    "武田信玄（軍師・取りまとめ）",
    "天下を束ねるがごとく、プロレス団体を一手に経営する。選手を育て、興行を打ち、ベルトを巡り覇権を争う——これは現代の戦国絵巻である。殿下のご意向のもと、我ら五名がこの要件を定める。",
    "1A0A00", "C8A96E", "C8A96E"
)

add_heading2(doc, "1.1 目的")
add_body(doc, "プレイヤーがプロレス団体のオーナーとなり、選手の獲得・育成・興行運営・タイトル管理を通じて団体を日本最高峰へと導くブラウザベースの経営シミュレーションゲームを開発する。")

add_heading2(doc, "1.2 対象ユーザー")
add_req_table(doc,
    ["ユーザー層", "特徴"],
    [
        ["プロレスファン",     "実際の団体運営に興味がある層。リアリティを重視"],
        ["経営シミュ好き",     "育成・数値管理・資源配分が好きなゲーマー"],
        ["カジュアルゲーマー", "難しすぎず、さくっと遊べることを期待"],
    ],
    col_widths=[5, 11]
)

add_heading2(doc, "1.3 技術スタック")
add_body(doc, "単一HTMLファイル（vanilla JS + CSS）として実装。サーバー・ビルドツール不要。ブラウザで直接開いて即プレイ可能。")

doc.add_page_break()

# ========== 2. 機能要件：経営・内政 ==========
add_heading1(doc, "2. 機能要件：経営・内政システム　【高坂昌信 担当】")
add_officer_comment(doc,
    "高坂昌信（逃げ弾正）",
    "軍略・内政・外交を一手に担う我が得意分野にございます。資金の流れを正しく設計せねば、いかに強い選手を揃えても団体は立ち行きませぬ。",
    "F0F0FD", "2A2A8B", "2A2A8B"
)

add_heading2(doc, "2.1 団体管理")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-101", "団体名・ロゴ設定", "ゲーム開始時に団体名を入力する",                   "高"],
        ["F-102", "資金管理",         "所持金の増減、収支履歴の表示",                     "高"],
        ["F-103", "ターン進行",       "1ターン＝1ヶ月。月末に固定費（給与等）を精算",     "高"],
        ["F-104", "ファン人気度",     "団体全体の人気値。興行収益・選手獲得交渉に影響",   "高"],
        ["F-105", "スポンサー契約",   "人気度に応じたスポンサー収入の自動計算",           "中"],
        ["F-106", "施設投資",         "道場・医療室などを建設し選手ステータスに補正",     "低"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

add_heading2(doc, "2.2 選手管理")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-201", "選手一覧",         "所属選手のステータス・契約状況を一覧表示",         "高"],
        ["F-202", "選手獲得（スカウト）", "資金を使いフリーや他団体の選手を獲得",         "高"],
        ["F-203", "選手育成",         "トレーニングにより能力値を成長させる",             "高"],
        ["F-204", "給与・契約管理",   "選手ごとの月俸設定。契約期間満了で離脱",           "中"],
        ["F-205", "コンディション管理","疲労・怪我状態を管理。無理使いで能力低下",        "中"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

add_heading3(doc, "選手ステータス定義")
add_req_table(doc,
    ["パラメータ", "説明", "範囲"],
    [
        ["パワー",       "投げ・打撃技の威力",           "1〜100"],
        ["スピード",     "スピード技・体力回復速度",     "1〜100"],
        ["テクニック",   "関節技・試合巧者度",           "1〜100"],
        ["タフネス",     "耐久力・試合体力",             "1〜100"],
        ["マイク",       "観客へのアピール力・人気への貢献", "1〜100"],
        ["コンディション","現在の体調（試合能力に補正）", "0〜100%"],
    ],
    col_widths=[3.5, 9, 3]
)

doc.add_page_break()

# ========== 3. 機能要件：興行・試合 ==========
add_heading1(doc, "3. 機能要件：興行・試合システム　【山県昌景 担当】")
add_officer_comment(doc,
    "山県昌景（赤備えの将）",
    "戦こそが全て！興行は我が担当。いかに強者を揃え、観衆を熱狂させるか——カードの組み方一つで興行の成否が決まりまする。先頭に立ち、最高の戦場を作ってみせましょう！",
    "FDF0F0", "8B0000", "8B0000"
)

add_heading2(doc, "3.1 興行開催")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-301", "興行計画",           "会場規模（小・中・大）と開催日を決定",             "高"],
        ["F-302", "カード編成",         "試合（3〜8試合）の対戦カードを組む",               "高"],
        ["F-303", "試合シミュレーション","ステータスをもとに勝敗・評価を自動計算",           "高"],
        ["F-304", "興行収益計算",       "観客動員数 × チケット単価で収益を算出",            "高"],
        ["F-305", "興行評価",           "試合内容・カードの組み合わせで評価点を算出",       "中"],
        ["F-306", "試合結果テキスト演出","試合経過を簡易テキスト実況で表示",                "中"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

add_heading2(doc, "3.2 試合シミュレーション詳細")
add_body(doc, "以下のフローで勝敗を決定する：")
add_body(doc, "選手ステータス取得　▶　コンディション補正　▶　乱数要素付加　▶　勝敗・評価算出　▶　結果反映（人気・疲労）")

add_heading2(doc, "3.3 タイトル（ベルト）管理")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-401", "タイトル定義", "団体内タイトルを最大3本設定可能",     "高"],
        ["F-402", "王者管理",     "現チャンピオンの記録・防衛回数管理", "高"],
        ["F-403", "挑戦者指名",   "次回防衛戦の挑戦者を指名できる",     "中"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

doc.add_page_break()

# ========== 4. ストーリー・抗争 ==========
add_heading1(doc, "4. 機能要件：ストーリー・抗争システム　【真田幸隆 担当】")
add_officer_comment(doc,
    "真田幸隆（調略の達人）",
    "兵を使わずして城を落とすのが我が流儀。プロレスもまた同じ——観客の心を動かすのは技の強さだけではなく、物語の力。ヒールとベビーフェイスの対立構図こそが、興行を何倍にも面白くする仕掛けにございます。",
    "FDFAF0", "8B6A00", "8B6A00"
)

add_heading2(doc, "4.1 ストーリーライン管理")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-501", "キャラ設定",         "各選手をヒール／ベビーフェイスに設定",                 "高"],
        ["F-502", "抗争設定",           "2選手間の抗争を登録。連続試合でヒートを上げる",       "中"],
        ["F-503", "ストーリーイベント", "裏切り・乱入などの特殊イベントを手動発動",             "低"],
        ["F-504", "抗争ヒート値",       "抗争の盛り上がり度を数値管理。興行評価に反映",         "中"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

add_heading2(doc, "4.2 他団体との関係")
add_req_table(doc,
    ["ID", "機能", "説明", "優先度"],
    [
        ["F-601", "他団体の存在", "CPU団体が複数存在し、業界内でランキング競争", "中"],
        ["F-602", "交流戦",       "他団体との対抗戦を開催し人気を獲得",         "低"],
    ],
    col_widths=[1.5, 4, 8, 2]
)

doc.add_page_break()

# ========== 5. 非機能要件 ==========
add_heading1(doc, "5. 非機能要件　【馬場信春 担当】")
add_officer_comment(doc,
    "馬場信春（不死身の鬼美濃）",
    "我が生涯、一度も傷を負わぬのは守りを疎かにせぬからにございます。いかに華やかな機能を揃えても、動かぬアプリでは意味がありませぬ。品質と堅牢さをこの馬場が保証いたします。",
    "F0FDF0", "2A6A2A", "2A6A2A"
)

add_req_table(doc,
    ["分類", "要件", "詳細"],
    [
        ["性能",  "ターン処理速度",   "1ターンの計算・描画を1秒以内に完了すること"],
        ["性能",  "試合シミュ速度",   "1興行（最大8試合）の計算を500ms以内に完了"],
        ["保存",  "セーブ機能",       "localStorage を使ったセーブ・ロード機能を実装"],
        ["保存",  "オートセーブ",     "各ターン終了時に自動保存"],
        ["UX",    "レスポンシブ対応", "PC・タブレット幅（768px以上）で正常表示"],
        ["UX",    "操作フィードバック","全てのボタン操作に即時の視覚的フィードバック"],
        ["互換性","ブラウザ",         "Chrome / Edge 最新版で動作すること"],
        ["構成",  "単一ファイル",     "HTML/CSS/JS を1ファイルに内包。外部依存なし"],
    ],
    col_widths=[2, 4.5, 9]
)

doc.add_page_break()

# ========== 6. 画面設計 ==========
add_heading1(doc, "6. 画面設計（概要）")
add_officer_comment(doc,
    "武田信玄",
    "高坂・山県・真田・馬場、各々の要件を束ね、画面の全体像を定める。シンプルでありながら全ての情報に素早くアクセスできる配置とする。",
    "1A0A00", "C8A96E", "C8A96E"
)

add_heading2(doc, "6.1 画面一覧")
add_req_table(doc,
    ["画面名", "主な内容"],
    [
        ["ダッシュボード（メイン）", "資金・人気・ターン数・直近イベントの概要表示"],
        ["選手管理画面",             "選手一覧・詳細・トレーニング・スカウト"],
        ["興行計画画面",             "カード編成・会場選択・開催ボタン"],
        ["興行結果画面",             "試合結果・収益・評価のサマリー"],
        ["タイトル管理画面",         "ベルト一覧・王者・挑戦者設定"],
        ["ストーリー画面",           "抗争管理・ヒール／ベビーフェイス設定"],
        ["設定画面",                 "セーブ・ロード・ゲームリセット"],
    ],
    col_widths=[6, 10]
)

add_heading2(doc, "6.2 ナビゲーション構造")
add_body(doc, "ダッシュボード　⇄　選手管理　⇄　興行計画　⇄　タイトル管理　⇄　ストーリー")
add_body(doc, "全画面は上部ナビゲーションバーから遷移可能。モーダルは使用せず、単一ページ内でのセクション切り替え方式を採用。")

doc.add_page_break()

# ========== 7. 開発フェーズ ==========
add_heading1(doc, "7. 開発フェーズ計画")
add_req_table(doc,
    ["フェーズ", "内容", "対象機能ID"],
    [
        ["Phase 1\nMVP",  "最小限で遊べる状態\nダッシュボード・選手管理・シンプル興行", "F-101〜104, F-201〜203\nF-301〜304, F-401〜402"],
        ["Phase 2\n拡充",  "ゲームとしての深みを追加\nコンディション・抗争・タイトル挑戦者", "F-205, F-305〜306\nF-403, F-501〜504"],
        ["Phase 3\n完成",  "経営要素の充実・他団体との競争", "F-105〜106, F-204\nF-601〜602"],
    ],
    col_widths=[3, 8, 5.5]
)

add_officer_comment(doc,
    "武田信玄（締めの言葉）",
    "疾きこと風の如く——まずはPhase 1で素早く動くものを作り、殿下にご覧いただく。徐として林の如く積み重ね、侵略すること火の如く拡大する。これが我らの策定した開発の道筋にございます。\n\n殿下のご下命を、謹んでお待ち申し上げております。",
    "1A0A00", "C8A96E", "C8A96E"
)

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_footer.add_run("仮想プロレス団体経営アプリ 要件定義書 v1.0　／　総大将：殿下　／　軍師：武田信玄\n配下：山県昌景・馬場信春・高坂昌信・真田幸隆")
set_font(r, size=8, color=(0xa0, 0x80, 0x60))

# ========== 保存 ==========
output_path = r"C:\Users\topge\OneDrive\ドキュメント\GitHub\claude_playground\docs\requirements_prowrestling.docx"
doc.save(output_path)
print(f"保存完了: {output_path}")
